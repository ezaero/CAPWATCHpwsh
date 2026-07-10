from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Cadet-Protection-M365-Controls.md"
OUTPUT = ROOT / "docs" / "Cadet-Protection-M365-Controls.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT_FILL = "F2F4F7"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_indent(table, dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_text = OxmlElement("w:t")
    fld_char_text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, fld_char_text, fld_char_end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run_element.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text, *, bold=False, color=INK, size=None):
    parts = re.split(r"(`[^`]+`|https?://\S+)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("http://") or part.startswith("https://"):
            add_hyperlink(paragraph, part, part)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(size if size else 10)
            run.font.color.rgb = color
            run.bold = bold
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = color
            if size:
                run.font.size = Pt(size)
            run.bold = bold


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.italic = True
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(2)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 6, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)


def add_running_furniture(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Civil Air Patrol Cadet Protection in Microsoft 365"
    header.style = doc.styles["Header"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_manual_numbered_paragraph(doc, marker, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    marker_run = paragraph.add_run(f"{marker}.\t")
    marker_run.font.color.rgb = INK
    add_inline_runs(paragraph, text)
    return paragraph


def add_table(doc, rows):
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_table_indent(table, 120)
    table_font_size = 9.5 if column_count >= 5 else 10.5
    margin_top_bottom = 60 if column_count >= 5 else 80

    width_map = {
        3: [2600, 2600, 4160],
        5: [2200, 1800, 2500, 1200, 1200],
    }
    widths = width_map.get(column_count, [int(9360 / column_count)] * column_count)

    header_cells = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        cell = header_cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell, top=margin_top_bottom, bottom=margin_top_bottom)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, text.strip(), bold=True, size=table_font_size)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for row_data in rows[1:]:
        row_cells = table.add_row().cells
        row_cells[0]._tc.getparent().get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for idx, text in enumerate(row_data):
            cell = row_cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, top=margin_top_bottom, bottom=margin_top_bottom)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, text.strip(), size=table_font_size)

    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_label_or_paragraph(doc, text):
    if text.endswith(":") and len(text) < 70:
        p = doc.add_paragraph()
        add_inline_runs(p, text, bold=True, color=DARK_BLUE)
    else:
        p = doc.add_paragraph()
        add_inline_runs(p, text)


def build_docx():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    add_running_furniture(doc)

    title_text = lines[0].lstrip("# ").strip()
    subtitle_text = lines[2].strip()
    prepared_text = lines[4].strip()

    title = doc.add_paragraph(style="Title")
    add_inline_runs(title, title_text, bold=True, color=RGBColor(0x0B, 0x25, 0x45))

    subtitle = doc.add_paragraph(style="Subtitle")
    add_inline_runs(subtitle, subtitle_text, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    add_inline_runs(meta, prepared_text, color=MUTED)

    i = 5
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3" if level == 4 else "Heading 1"
            p = doc.add_paragraph(style=style)
            add_inline_runs(p, text, bold=True, color=BLUE if style != "Heading 3" else DARK_BLUE)
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            add_manual_numbered_paragraph(doc, numbered.group(1), numbered.group(2))
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            i += 1
            continue

        add_label_or_paragraph(doc, line)
        i += 1

    doc.core_properties.title = title_text
    doc.core_properties.subject = "Working group discussion brief"
    doc.core_properties.author = "CAPWATCHpwsh"
    doc.core_properties.keywords = "Civil Air Patrol, Cadet Protection, Microsoft 365, Exchange Online, Teams, Purview"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
