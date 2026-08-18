import csv
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "office365-distribution-lists-2026-08-12.csv"
OUT_PATH = ROOT / "docs" / "CAPNET-Distribution-Groups-Teams-Mail-Enabled-Security-Groups-Policy.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"


SPECIALTY_TRACKS = [
    "Administration",
    "Aerospace",
    "Cadet Programs",
    "Chaplain",
    "Character Development",
    "Communications",
    "Drug Demand Reduction",
    "Emergency Services",
    "Finance",
    "Flight Operations",
    "Health Services",
    "Historian",
    "Information Technology Officer",
    "Inspector General",
    "Legal",
    "Logistics",
    "Operations",
    "Personnel",
    "Plans And Programs",
    "Professional Development",
    "Public Affairs",
    "Recruiting And Retention Officer",
    "Safety",
    "Standards And Evaluations",
]

OPS_QUAL_LISTS = [
    "Aircrew",
    "CO Wing Check Pilots",
    "CO Wing ESOfficers",
    "CO WING INCIDENT COMMANDERS",
    "CO Wing Mission Flying",
    "CO Wing Pilot List",
    "CO Wing Staff",
    "CO Wing Stan Eval",
    "Communicators",
    "ESList",
    "Group 4 Emergency Services Officers",
    "Instructor Pilots",
    "KAPA Pilot List",
    "Mission Base Staff",
    "Mission Check Pilots",
    "Mission Pilots",
    "Orientation Pilots",
    "Pilots",
    "sUAS",
]


def load_rows():
    with CSV_PATH.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths):
    table.style = "Table Grid"
    table.autofit = False
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, color=BLACK)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run_font(run, size=9.5, color=BLACK)


def set_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("CAPNET Distribution Groups, Teams, and Mail-Enabled Security Groups")
    set_run_font(run, size=9, color=GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    run = footer.add_run("Civil Air Patrol | CAPNET Governance")
    set_run_font(run, size=9, color=GRAY)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("12 August 2026")
    set_run_font(run, size=11, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CAPNET")
    set_run_font(run, size=24, color=BLACK, bold=True)

    for text in [
        "Distribution Groups, Teams,",
        "and Mail-Enabled Security Groups",
        "Governance Policy",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=22, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("NATIONAL HEADQUARTERS CIVIL AIR PATROL")
    set_run_font(run, size=11, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Maxwell Air Force Base, Alabama")
    set_run_font(run, size=11, color=BLACK)


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, data, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in data:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = "" if value is None else str(value)
    style_table(table, widths)
    return table


def row_for(rows, display_name, category=None):
    matches = [r for r in rows if r["DisplayName"] == display_name and (category is None or r["Category"] == category)]
    return sorted(matches, key=lambda r: (r["Category"], r["PrimarySmtpAddress"]))


def make_doc():
    rows = load_rows()
    counts = Counter(r["Category"] for r in rows)

    doc = Document()
    set_styles(doc)
    set_header_footer(doc)
    add_title_block(doc)

    add_para(
        doc,
        "This document establishes the national CAPNET standard for distribution lists, dynamic "
        "distribution lists, mail-enabled security groups, Microsoft 365 groups, and Teams. Local "
        "wing examples are included only to illustrate how the standard is applied.",
    )

    doc.add_heading("Background", level=1)
    add_para(
        doc,
        "Civil Air Patrol is consolidating identity, email, collaboration, and communications into "
        "CAPNET. Because CAP includes National Headquarters, regions, wings, groups, squadrons, "
        "flights, schools, activities, and mission programs, communication objects must be named "
        "with a consistent scope prefix. That prefix allows CAPWATCH-driven automation and approved "
        "CAPNET administration processes to maintain membership without local manual drift or "
        "ambiguous ownership.",
    )

    doc.add_heading("Objective", level=2)
    add_para(
        doc,
        "Define the national naming, address, ownership, and lifecycle standard for CAPNET "
        "communication objects. The standard must work for every wing by substituting the approved "
        "wing designator, unit charter number, audience name, and CAPNET mail domain into the same "
        "patterns.",
    )

    doc.add_heading("Policy", level=1)
    add_bullets(
        doc,
        [
            "CAPNET communication objects shall use a scope prefix for National Headquarters, region, wing, group, unit, activity, or program ownership.",
            "Wing-scoped objects shall use the approved two-letter wing designator in display names and SMTP local parts. For example, Colorado Wing uses CO, Texas Wing uses TX, and California Wing uses CA.",
            "Unit-scoped objects shall use the wing designator and three-digit unit charter number in the form <Wing>-<Unit>, such as CO-015, TX-001, or CA-123.",
            "CAPWATCH or another approved authoritative source shall drive membership for patterned groups when the source data can reliably identify the audience.",
            "Teams and Microsoft 365 groups shall be created only when collaboration artifacts are required. Email-only audiences shall remain distribution lists, dynamic distribution lists, or mail-enabled security groups.",
            "Mail-enabled security groups shall be used when the object must also grant access, scope permissions, or support an externally reachable role mailbox pattern.",
            "Any exception to the national pattern shall have a named owner, business purpose, migration rationale, and periodic review date.",
        ],
    )

    doc.add_heading("Scope and Wing Designators", level=1)
    add_para(
        doc,
        "The scope token identifies who owns the object and which authoritative data source should "
        "drive membership. Display names should use readable uppercase designators. SMTP local parts "
        "should use lowercase letters, numbers, and hyphens. The domain token <capnet-domain> means "
        "the approved CAPNET mail domain for that object, such as the default member domain or an "
        "approved mission domain.",
    )
    scope_rows = [
        ["National Headquarters", "NHQ", "NHQ <Audience or Function>", "nhq-<audience>@<capnet-domain>"],
        ["Region", "<Region>", "<Region> <Audience or Function>", "<region>-<audience>@<capnet-domain>"],
        ["Wing", "<Wing>", "<Wing> Wing <Audience or Function>", "<wing>-<audience>@<capnet-domain>"],
        ["Wing group", "<Wing> Group <N>", "<Wing> Group <N> <Audience>", "<wing>-group-<n>-<audience>@<capnet-domain>"],
        ["Unit", "<Wing>-<Unit>", "<Wing>-<Unit> <Audience or Unit Name>", "<wing>-<unit>-<audience>@<capnet-domain>"],
        ["Activity or school", "<Scope>-<Activity>", "<Scope> <Activity> <Audience>", "<scope>-<activity>-<audience>@<capnet-domain>"],
    ]
    add_table(doc, ["Scope", "Token", "Display-name standard", "SMTP local-part standard"], scope_rows, [1850, 1700, 2850, 2960])

    doc.add_heading("Wing Examples", level=2)
    wing_examples = [
        ["Colorado Wing", "CO", "CO-015 Cadets", "co-015-cadets@<capnet-domain>"],
        ["Texas Wing", "TX", "TX-001 Seniors", "tx-001-seniors@<capnet-domain>"],
        ["California Wing", "CA", "CA-123 Recruiting", "ca-123-recruiting@<capnet-domain>"],
        ["Florida Wing", "FL", "FL Wing Announcements", "fl-announcements@<capnet-domain>"],
        ["Alaska Wing", "AK", "AK Group 1 Commanders", "ak-group-1-commanders@<capnet-domain>"],
    ]
    add_table(doc, ["Example wing", "Designator", "Example display name", "Example SMTP address"], wing_examples, [1850, 1350, 2850, 3310])

    doc.add_page_break()
    doc.add_heading("National Object Patterns", level=1)
    pattern_rows = [
        [
            "Wing-wide communications",
            "<Wing> Wing <Audience or Function>",
            "<wing>-<audience>@<capnet-domain>",
            "Wing-level announcements, staff, senior, cadet, and functional-office lists.",
        ],
        [
            "Unit all-member list",
            "<Wing>-<Unit> <Unit Name>",
            "<wing>-<unit>@<capnet-domain>",
            "One list per unit for all members assigned to the unit.",
        ],
        [
            "Unit cadet list",
            "<Wing>-<Unit> Cadets",
            "<wing>-<unit>-cadets@<capnet-domain>",
            "Cadets plus approved parent and cadet-program staff inclusion rules.",
        ],
        [
            "Unit senior list",
            "<Wing>-<Unit> Seniors",
            "<wing>-<unit>-seniors@<capnet-domain>",
            "Senior members assigned to the unit.",
        ],
        [
            "Recruiting group",
            "<Wing>-<Unit> Recruiting",
            "<wing>-<unit>-recruiting@<capnet-domain>",
            "Mail-enabled security group for unit recruiting, retention, command, and executive staff routing.",
        ],
        [
            "Wing group list",
            "<Wing> Group <N> <Audience>",
            "<wing>-group-<n>-<audience>@<capnet-domain>",
            "Sub-wing geographic or command-group distribution based on predefined unit membership.",
        ],
        [
            "Region list",
            "<Region> <Audience or Function>",
            "<region>-<audience>@<capnet-domain>",
            "Region-level audiences driven by region staff, wing command roles, or activity assignment.",
        ],
        [
            "Specialty-track list",
            "<Scope> Specialty Track <Track>",
            "<scope>-spec-<track-slug>@<capnet-domain>",
            "Audience is defined by the authoritative specialty-track title and member status.",
        ],
        [
            "Operations qualification list",
            "<Scope> Ops <Qualification>",
            "<scope>-ops-<qualification-slug>@<capnet-domain>",
            "Audience is defined by approved operational qualification, task, or achievement data.",
        ],
        [
            "Unit Team",
            "<Wing>-<Unit> <Unit or Team Name>",
            "<wing>-<unit>-team@<capnet-domain>",
            "Microsoft Team for collaboration within a unit or staff section.",
        ],
        [
            "Wing functional Team",
            "<Wing> <Function>",
            "<wing>-<function>-team@<capnet-domain>",
            "Microsoft Team for wing staff collaboration, project work, or program operations.",
        ],
    ]
    add_table(doc, ["Family", "Display-name pattern", "SMTP pattern", "Use"], pattern_rows, [1850, 2300, 2300, 2910])

    doc.add_heading("Display Name and Address Rules", level=1)
    add_bullets(
        doc,
        [
            "Display names shall be readable and shall begin with the scope token unless the object is a national object whose ownership is evident from the NHQ prefix.",
            "SMTP local parts shall be lowercase and hyphenated. Spaces, ampersands, underscores, and local nicknames shall not be used in new primary addresses.",
            "The primary SMTP address shall be canonical. Legacy wing or local addresses may remain as proxy addresses when required for continuity.",
            "The same pattern shall be used for every wing. The wing designator changes, not the structure of the object name.",
            "Generated aliases shall not omit the scope token. For example, use co-emergency-services or tx-emergency-services, not emergencyservices, when operating in a national tenant.",
        ],
    )

    doc.add_heading("Specialty Track Lists", level=1)
    add_para(
        doc,
        "Specialty-track lists shall be scoped nationally, regionally, or by wing. The CAPWATCH specialty-track title remains the authoritative business label, but the object name must include the owning scope so that the same specialty track can exist in multiple wings without collision.",
    )
    specialty_pattern_data = [
        ["Wing specialty track", "<Wing> Specialty Track <Track>", "<wing>-spec-<track-slug>@<capnet-domain>"],
        ["Region specialty track", "<Region> Specialty Track <Track>", "<region>-spec-<track-slug>@<capnet-domain>"],
        ["National specialty track", "NHQ Specialty Track <Track>", "nhq-spec-<track-slug>@<capnet-domain>"],
    ]
    add_table(doc, ["Scope", "Display-name standard", "SMTP standard"], specialty_pattern_data, [2500, 3650, 3050])
    add_para(
        doc,
        "Colorado implementation examples from the current export are listed below to show the source program names that should be slugged and scoped in CAPNET.",
    )
    specialty_data = []
    for name in SPECIALTY_TRACKS:
        for r in row_for(rows, name, "DistributionList"):
            specialty_data.append([r["DisplayName"], r["PrimarySmtpAddress"], r["Category"]])
    add_table(doc, ["Current example", "Current primary SMTP address", "Object family"], specialty_data, [3600, 3700, 1900])

    doc.add_heading("Operations Qualification Lists", level=1)
    add_para(
        doc,
        "Operations-qualification lists shall be scoped to the operational authority that uses the audience. "
        "The qualification or role name should be written plainly in the display name and slugged in the SMTP local part. "
        "Local airport, aircraft, exercise, or incident names may be used only after the wing or activity scope token.",
    )
    ops_pattern_data = [
        ["Wing qualification", "<Wing> Ops <Qualification>", "<wing>-ops-<qualification-slug>@<capnet-domain>"],
        ["Wing group qualification", "<Wing> Group <N> Ops <Qualification>", "<wing>-group-<n>-ops-<qualification-slug>@<capnet-domain>"],
        ["Activity or mission qualification", "<Scope> <Activity> Ops <Qualification>", "<scope>-<activity>-ops-<qualification-slug>@<capnet-domain>"],
        ["National qualification", "NHQ Ops <Qualification>", "nhq-ops-<qualification-slug>@<capnet-domain>"],
    ]
    add_table(doc, ["Scope", "Display-name standard", "SMTP standard"], ops_pattern_data, [2500, 3650, 3050])
    add_para(
        doc,
        "Colorado implementation examples from the current export are listed below because the existing names mix qualifications, role labels, and local operating areas.",
    )
    ops_data = []
    for name in OPS_QUAL_LISTS:
        for r in row_for(rows, name):
            ops_data.append([r["DisplayName"], r["PrimarySmtpAddress"], r["Category"]])
    add_table(doc, ["Current example", "Current primary SMTP address", "Object family"], ops_data, [3600, 3700, 1900])

    doc.add_heading("Mail-Enabled Security Groups", level=1)
    add_para(
        doc,
        "Mail-enabled security groups shall be used when the object both distributes mail and scopes authorization, application access, calendar access, or another security boundary. "
        "Recruiting groups are the standard unit-level mail-enabled security group because they commonly require external inquiry routing and internal role-based access control.",
    )
    mesg_pattern_data = [
        ["Unit recruiting", "<Wing>-<Unit> Recruiting", "<wing>-<unit>-recruiting@<capnet-domain>"],
        ["Scoped application access", "<Scope> Security <Application or Permission>", "<scope>-sec-<permission>@<capnet-domain>"],
        ["Calendar or resource editors", "<Scope> <Resource> Editors", "<scope>-<resource>-editors@<capnet-domain>"],
    ]
    add_table(doc, ["Use", "Display-name standard", "SMTP standard"], mesg_pattern_data, [2600, 3500, 3100])
    add_para(
        doc,
        "Non-recruiting mail-enabled security groups from the Colorado export are listed below as migration review examples.",
    )
    non_recruiting_mesg = [
        r
        for r in rows
        if r["Category"] == "MailEnabledSecurityGroup"
        and not (r["DisplayName"].startswith("CO-") and r["DisplayName"].endswith(" Recruiting"))
    ]
    add_table(
        doc,
        ["Non-recruiting mail-enabled security group", "Primary SMTP address"],
        [[r["DisplayName"], r["PrimarySmtpAddress"]] for r in sorted(non_recruiting_mesg, key=lambda r: r["DisplayName"].lower())],
        [4550, 4650],
    )

    doc.add_heading("Microsoft 365 Groups and Teams", level=1)
    add_para(
        doc,
        "Microsoft 365 groups without Teams are mailbox-backed collaboration groups. Teams are Microsoft 365 groups with the Team workload provisioned. "
        "CAPNET shall use Teams for collaboration spaces that need channels, files, meetings, or persistent chat. Teams shall not be created solely to obtain an email address.",
    )
    team_pattern_data = [
        ["Unit Team", "<Wing>-<Unit> <Unit Name>", "<wing>-<unit>-team@<capnet-domain>"],
        ["Wing staff Team", "<Wing> Wing Staff", "<wing>-wing-staff-team@<capnet-domain>"],
        ["Wing functional Team", "<Wing> <Function>", "<wing>-<function>-team@<capnet-domain>"],
        ["Region Team", "<Region> <Function>", "<region>-<function>-team@<capnet-domain>"],
        ["National Team", "NHQ <Function>", "nhq-<function>-team@<capnet-domain>"],
    ]
    add_table(doc, ["Use", "Display-name standard", "SMTP standard"], team_pattern_data, [2600, 3500, 3100])
    add_para(
        doc,
        "Microsoft 365 groups without Teams from the Colorado export are listed below as examples requiring business-owner validation.",
    )
    m365 = [r for r in rows if r["Category"] == "Microsoft365Group"]
    add_table(
        doc,
        ["Microsoft 365 group", "Primary SMTP address"],
        [[r["DisplayName"], r["PrimarySmtpAddress"]] for r in sorted(m365, key=lambda r: r["DisplayName"].lower())],
        [4550, 4650],
    )

    doc.add_heading("Dynamic Distribution Groups", level=1)
    add_para(
        doc,
        "Dynamic distribution groups shall be used only when the recipient filter can be expressed using CAPNET-supported attributes such as unit assignment, wing assignment, duty position, staff assignment, or member category. "
        "Each dynamic list shall include the scope token in both display name and SMTP address.",
    )
    dynamic_pattern_data = [
        ["Wing staff", "<Wing> Wing Staff", "<wing>-wing-staff@<capnet-domain>"],
        ["Wing commanders", "<Wing> Commanders", "<wing>-commanders@<capnet-domain>"],
        ["Wing group commanders", "<Wing> Group <N> Commanders", "<wing>-group-<n>-commanders@<capnet-domain>"],
        ["Region directors", "<Region> Directors", "<region>-directors@<capnet-domain>"],
    ]
    add_table(doc, ["Use", "Display-name standard", "SMTP standard"], dynamic_pattern_data, [2600, 3500, 3100])
    add_para(
        doc,
        "Dynamic distribution groups from the Colorado export are listed below as examples requiring filter review before CAPNET migration.",
    )
    dynamic_groups = [r for r in rows if r["Category"] == "DynamicDistributionList"]
    add_table(
        doc,
        ["Dynamic distribution group", "Primary SMTP address"],
        [[r["DisplayName"], r["PrimarySmtpAddress"]] for r in sorted(dynamic_groups, key=lambda r: r["DisplayName"].lower())],
        [4550, 4650],
    )

    doc.add_heading("Governance Requirements", level=1)
    add_bullets(
        doc,
        [
            "Each patterned family shall have an approved provisioning rule, naming rule, authoritative data source, owner role, and retirement rule.",
            "Each object shall have an owner at the same or higher scope as the object. A wing object is owned by a wing role; a region object is owned by a region role; a national object is owned by NHQ.",
            "Each unpatterned list, security group, Microsoft 365 group, or Team shall be mapped to a business owner before migration or disabled when no owner can be identified.",
            "Aliases shall be normalized to lowercase, hyphenated, scoped local parts for generated groups unless a legacy address must remain as a proxy address.",
            "Teams shall not be treated as mail-enabled security boundaries. Where a Team also needs access control, the access policy shall be documented separately.",
            "Recruiting groups should allow external senders only where the mission requires public inquiry routing and the group has an accountable owner.",
            "CAPNET production automation should run in preview mode before enforcing list membership changes at national scale.",
        ],
    )

    doc.add_heading("Migration from Local Wing Patterns", level=1)
    add_para(
        doc,
        "Existing local wing objects should be transformed by replacing local names, local domains, and implicit wing context with explicit CAPNET scope tokens. "
        "For example, a local Colorado object named Emergency Services should become CO Specialty Track Emergency Services or CO Ops Emergency Services depending on its source rule. "
        "A local object named CO-015 Cadets should retain the display-name structure but use the national primary SMTP pattern co-015-cadets@<capnet-domain>. "
        "The national tenant should preserve required legacy proxy addresses during transition while enforcing one canonical primary SMTP address per object.",
    )

    doc.add_heading("Colorado Export as Implementation Evidence", level=1)
    summary_data = [
        [category, counts[category]]
        for category in [
            "DistributionList",
            "DynamicDistributionList",
            "MailEnabledSecurityGroup",
            "Microsoft365Group",
            "Team",
        ]
    ]
    add_table(doc, ["Object family", "Count in 12 Aug 2026 Colorado export"], summary_data, [6200, 3000])
    unit_counts = {
        "Unit all-member lists": sum(
            1
            for r in rows
            if r["DisplayName"].startswith("CO-")
            and r["PrimarySmtpAddress"].lower().startswith(r["DisplayName"][:6].lower() + "@")
        ),
        "Unit cadet lists": sum(1 for r in rows if r["DisplayName"].startswith("CO-") and r["DisplayName"].endswith(" Cadets")),
        "Unit senior lists": sum(1 for r in rows if r["DisplayName"].startswith("CO-") and r["DisplayName"].endswith(" Seniors")),
        "Recruiting mail-enabled security groups": sum(1 for r in rows if r["DisplayName"].startswith("CO-") and r["DisplayName"].endswith(" Recruiting")),
        "Unit Teams": sum(1 for r in rows if r["Category"] == "Team" and r["DisplayName"].startswith("CO-")),
    }
    add_table(doc, ["Observed Colorado family", "Observed count"], [[k, v] for k, v in unit_counts.items()], [6500, 2700])
    add_para(
        doc,
        f"Source inventory: {CSV_PATH.name}. Total Colorado objects reviewed: {len(rows)}. "
        "These counts are not national requirements; they are implementation evidence used to validate that the national standard covers existing wing patterns.",
    )

    doc.add_heading("Appendix A: Source and Review Notes", level=1)
    add_para(
        doc,
        "The implementation evidence was exported from Exchange Online on 12 August 2026 and reviewed alongside the CAPWATCH PowerShell automation scripts. "
        "The export included classic distribution groups, dynamic distribution groups, mail-enabled security groups, Microsoft 365 groups, and Teams-backed groups. "
        "The CAPNET policy format was modeled on the sibling CAPNET architecture, licensing, and automation overview documents, using the same governance-document shape of Background, Objective, policy sections, and appendices.",
    )

    doc.core_properties.title = "CAPNET Distribution Groups, Teams, and Mail-Enabled Security Groups Governance Policy"
    doc.core_properties.subject = "CAPNET M365 communication group governance"
    doc.core_properties.author = "CAP M365 Working Group"
    doc.core_properties.comments = "Generated from Office 365 group inventory export and CAPWATCH automation scripts; revised as a CAPNET national standard."

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(make_doc())
